# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width  = Emu(7560310)
    section.page_height = Emu(10692130)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# Helper functions
def add_para(text, font_size=16, bold=False, color=None, alignment=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_field(label, value, font_size=16):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.font.size = Pt(font_size)
    r2 = p.add_run(value)
    r2.font.size = Pt(font_size)
    r2.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    return p

def add_heading_text(text, font_size=16):
    p = doc.add_paragraph()
    p.style = doc.styles['List Paragraph']
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    return p

def add_bullet_item(text, font_size=11, color=(0xEE, 0x00, 0x00)):
    p = doc.add_paragraph()
    p.style = doc.styles['List Paragraph']
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.bold = True
    r.font.color.rgb = RGBColor(*color)
    return p

def add_innov_item(text, font_size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.bold = True
    return p

def add_tech_subtitle(text, font_size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    r.bold = True
    return p

def add_tech_body(text, font_size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(font_size)
    return p

def add_dash_item(prefix, body, font_size=12):
    p = doc.add_paragraph()
    r1 = p.add_run(prefix)
    r1.font.size = Pt(font_size)
    r1.bold = True
    r2 = p.add_run(body)
    r2.font.size = Pt(font_size)
    return p

# ===== TITLE =====
add_para('文献分享核心信息', font_size=18)

# ===== META INFO =====
add_field('标题：', '      Do As I Can, Not As I Say: Grounding Language in Robotic Affordances')
add_field('期刊：', '      CoRL 2022 (Conference on Robot Learning)')
add_field('作者：', '      Michael Ahn, Anthony Brohan, Noah Brown, ... (Google Robotics & Everyday Robots)')
add_field('分享人：', '    张子杰')
add_field('日期：', '      2026-6-9')

add_para('')

# ===== 一、研究背景 =====
add_heading_text('一、研究背景', font_size=16)

add_bullet_item('1. 大语言模型（LLM）的"知行脱节"困境：LLM 在训练中编码了丰富的语义知识，但它们从未与物理世界发生过交互——既没有观察过自己生成的话语会对物理过程产生什么后果，也完全不了解一个具体机器人 agent 在当前环境中能做哪些事、不能做哪些事。例如，对一个只能"拿海绵"和"开关抽屉"的厨房机器人说"我打翻了饮料，能帮我清理吗？"，纯 LLM 可能回答"用吸尘器吸干净"——但机器人根本没有吸尘器，也不会使用吸尘器。')

add_bullet_item('2. LLM 不了解当前环境的实际状态：即使 LLM 知道机器人有哪些技能，它也无法感知当前场景的具体情况。比如指令"给我拿一个苹果"，但如果厨房里根本没有苹果，或者苹果在关着的抽屉里需要先打开抽屉才能拿到，纯 LLM 是无法知道这些的。机器人手里已经拿着一个苹果了，再让它"拿一个苹果"也不合理。LLM 缺乏这种对物理状态的实时感知能力。')

add_bullet_item('3. 现有方法的根本缺陷：此前的方法要么让 LLM 直接生成任务计划然后交给机器人执行——计划是在"真空"中产生的，完全没有考虑机器人的能力约束和当前环境状态；要么尝试通过微调让 LLM 接受视觉或状态输入——需要大量额外训练，泛化能力有限。核心问题在于：如何提取 LLM 中的语义知识来指导机器人执行真实世界中的长时域、抽象的自然语言指令，同时确保每一步决策都既对任务有用（useful）、又在当前状态下可行（feasible）？')

add_para('')
add_para('论文创新和贡献点：', font_size=16)

add_innov_item('1. 首次提出 SayCan 框架，将 LLM 的语义知识与机器人技能的 Affordance（价值函数）进行概率耦合，实现"任务接地"与"世界接地"的统一。')
add_innov_item('2. 提出了一种零样本、无需额外训练的长时域任务执行范式——通过 prompt engineering 实现泛化，不依赖任何针对新任务的微调。')
add_innov_item('3. 在真实世界进行了系统级的验证：使用 Everyday Robots 移动操作机器人，在真实厨房环境下对 101 条指令（7 个类别）进行了全面评估，同时在全新厨房中测试了泛化能力。')
add_innov_item('4. 首次实证了 NLP 的进步可以直接转化为机器人表现的提升：对比 PaLM 8B→62B→540B 及 FLAN 137B，机器人执行成功率随 LLM 能力增强而单调递增，错误率减少近一半。')
add_innov_item('5. 开源了基于表格环境的 SayCan 实现，降低了复现门槛。')

add_para('')

# ===== 二、关键技术（方法设计） =====
add_heading_text('二、关键技术（方法设计）', font_size=16)

add_para('')

add_tech_subtitle('1. 核心公式：概率乘积的直觉', font_size=12)
add_tech_body('SayCan 的核心可归结为一个极其简洁的概率乘积公式：', font_size=12)
add_tech_body('    P(技能能推进指令) ∝ P(技能在当前状态下能成功) × P(该技能对完成指令是合适的)', font_size=12)
add_tech_body('前者由 RL 学到的价值函数（Value Function）提供，论文称之为"世界接地"（world-grounding）或"Can"；后者由 LLM 对技能文本描述的评分提供，论文称之为"任务接地"（task-grounding）或"Say"。两个来源不同的概率信号相乘，用一个简单的乘积实现了"天"（LLM 的语义知识）与"地"（机器人的物理经验）的统一。', font_size=12)
add_tech_body('直觉：如果一个技能在当前状态下不可能成功（如"拿起可乐罐"但罐子不在视野里），那么无论 LLM 给它多高的分数，联合概率都应该很低；如果一个技能 LLM 认为对指令毫无帮助，那么即使机器人能完美执行它，联合概率也应该很低。只有两个条件同时满足，该技能才会被选中。这就是 SayCan 的精髓：LLM 动口，机器人动手，但只有"动手能做到"的口头建议才会被执行。', font_size=12)

add_para('')

add_tech_subtitle('2. 双重组件的详解', font_size=12)

add_dash_item('---LLM 的语言评分（"Say"部分）：', '将规划过程构建为"用户-机器人"的对话形式，通过精心设计的 prompt，引导 LLM 以结构化格式输出并利用 scoring mode 比较所有候选技能文本描述在当前上下文中的 log-likelihood，从而获得每个技能的定量评分。', font_size=12)

add_dash_item('---Value Function 的 Affordance 评分（"Can"部分）：', '关键洞察：在稀疏奖励（成功=1，失败=0）且不考虑折扣因子的强化学习设定中，状态-动作价值函数 Q(s,a) 天然等价于一个 affordance 函数——表示"在状态 s 下执行技能 a，最终能成功的概率"。利用 MT-Opt 框架在仿真中并行训练数百个语言条件化的技能，每个技能对应一个能从图像观测直接输出 Q 值的神经网络。当机器人面对新状态时，所有技能的 Q 值一次性被计算，形成一个"affordance 空间"——等价于机器人在心里对所有技能做了快速扫描："拿起可乐罐？0.95（视野中有可乐）。拿起苹果？0.02（没有苹果）。"', font_size=12)

add_para('')

add_tech_subtitle('3. 算法流程（迭代式决策循环）', font_size=12)
add_tech_body('SayCan 的运行是一个"边看边想边做"的迭代循环：接收用户指令和当前相机观测 → 将所有 N 个候选技能的文本描述输入 LLM，得到 p_LLM → 将当前图像观测输入所有技能的价值函数网络，得到 p_afford → 计算联合概率 p_combined = p_afford × p_LLM，选最高分技能 → 执行该技能，更新状态 → 将已执行的技能追加到 LLM 上下文中，回到第二步，直到选中终止标记"done"。', font_size=12)

add_para('')

add_tech_subtitle('4. 低层技能的获取', font_size=12)
add_tech_body('论文使用了两种互补的训练方式：（1）模仿学习（Behavioral Cloning, BC）：从人类遥操作收集的大量演示数据中学习策略；（2）强化学习（RL）：在仿真环境中使用 MT-Opt 框架训练语言条件化的价值函数和策略，并用 RetinaGAN 做 sim-to-real 风格迁移。奖励设置为稀疏的：episode 结束时由 3 人评审（2/3 同意即为成功），成功=1，失败=0。论文共定义了 551 个技能，分为 Pick、Place、Open/Close Drawer、Navigate to、Bring to 等技能族。', font_size=12)

add_para('')

add_tech_subtitle('5. 扩展能力', font_size=12)
add_tech_body('集成 Chain-of-Thought prompting 后，SayCan 能处理需要推理的指令（如"给我一杯不含咖啡因的水果味饮料"，模型先推理出"柠檬汽水符合条件"再选技能）；由于 PaLM 训练语料包含多语言数据，SayCan 天然支持非英语的自然语言指令。', font_size=12)

add_para('')

# ===== 三、实验结果分析 =====
add_heading_text('三、实验结果分析', font_size=16)

add_para('')
add_tech_subtitle('实验规模', font_size=12)
add_tech_body('使用 Everyday Robots 移动操作机器人（7-DoF 手臂 + 两指抓手），在真实办公室厨房和模拟厨房中进行评估。共设计 101 条测试指令，分为 7 个类别：NL Single Primitive（单步原语）、NL Nouns（名词推理）、NL Verbs（动词推理）、Structured（结构化消融）、Embodiment（具身化变体）、Crowd Sourced（众包自然语言）、Long-Horizon（10+ 步长时域任务）。评估指标包括规划成功率（3 人评判计划是否正确）和执行成功率（3 人看执行视频评判是否完成任务）。', font_size=12)

add_para('')
add_tech_subtitle('核心实验结果', font_size=12)

add_dash_item('---完整方法 vs 消融对比：', 'PaLM-SayCan（完整方法）在训练厨房中取得了 84% 的规划成功率和 74% 的执行成功率；在全新真实厨房中规划 81%、执行 60%——具有不错的泛化能力。去掉 Value Function（仅靠 LLM 评分选技能，无 affordance 约束），规划成功率骤降至 67%——LLM 常选语义合理但当前状态下不可执行的技能。去掉 LLM（直接使用 BC 策略或通过 USE 嵌入做语义匹配），执行成功率几乎为 0%——BC 策略完全不理解高层指令的语义。两个组件的必要性得到了双向验证：没有 LLM，机器人不知道"该做什么"；没有 affordance，机器人不知道"能做什么"。', font_size=12)

add_dash_item('---LLM 规模的影响（首创性发现）：', '对比 PaLM 8B、62B、540B 以及 FLAN 137B：PaLM 540B 比 FLAN 137B 在整个执行成功率上高出 14 个百分点（74% vs 60%），将错误率减少了近一半。LLM 越大，机器人越强——这意味着 NLP 领域每一次基础模型升级，都可以直接转化为机器人表现的提升，无需改动机器人系统本身。', font_size=12)

add_dash_item('---最具代表性的案例：', '指令"我刚锻炼完，能帮我拿些饮料和零食来恢复吗？"→ SayCan 执行：找水瓶→拿水瓶→带给你→放下→找苹果→拿苹果→带给你→完成。展现了 LLM 理解"恢复锻炼"隐含语义（需要健康食品）、自主确定需要几件物品、并在执行过程中正确维持长时域任务记忆的能力。指令"我留了可乐、苹果和水在外面，把它们扔掉然后拿海绵给我擦桌子"→ SayCan 正确理解"它们"指代三件物品、垃圾要去垃圾桶、擦桌子需要海绵，并成功执行了这一复杂序列。', font_size=12)

add_para('')
add_tech_subtitle('失败分析', font_size=12)
add_tech_body('在整体失败案例中，65% 归因于 LLM 错误（早期终止、否定/歧义处理不足），35% 归因于 affordance 错误（Value Function 对状态判断偏差）。低层技能的物理执行本身也存在失败率（抓取失败、碰撞等），进一步拉低了端到端成功率。技能库的覆盖范围和鲁棒性仍然是整个系统的核心瓶颈。', font_size=12)

add_para('')

# ===== 四、总结与展望 =====
add_heading_text('四、总结与展望', font_size=16)

add_para('')
add_tech_subtitle('论文核心总结', font_size=12)
add_tech_body('SayCan 提出了一个简洁而强大的方法论框架，将 LLM 中丰富的语义知识"接地"到真实机器人的物理行动中。核心洞察：LLM 擅长"说"（知道该做什么），但只有通过 RL 学到的 affordance 才知道"能做"（在当前状态下能成功做什么），将两者以概率乘积的方式结合起来，就能让机器人既理解高层指令的意图，又尊重自身能力和环境条件的约束。论文标题"Do As I Can, Not As I Say"一语中的：机器人只做它"能做"的事，而不是 LLM"说"该做的事。affordance 函数是那个关键的过滤器——确保每条来自 LLM 的建议都必须经过"可行性审查"，只有真正在当前场景中可执行的技能才会被选中。三层能力解耦（高层语义推理→LLM / 可行性判断→Value Function / 运动控制→BC/RL 策略）使得每一层都可以独立升级。', font_size=12)

add_para('')
add_tech_subtitle('局限性', font_size=12)
add_dash_item('---语言层面的局限性继承：', 'SayCan 继承 LLM 的所有已知弱点，包括对否定、量词、歧义和复杂推理的处理不足，以及对训练数据中社会偏见的继承。', font_size=12)
add_dash_item('---技能库是系统上限：', '无论 LLM 多么聪明，SayCan 只能组合现有技能。技能的数量、质量和鲁棒性直接决定整个系统的能力天花板。', font_size=12)
add_dash_item('---开环规划缺乏错误恢复：', 'SayCan 按既定序列执行，如果某个技能意外失败（尽管 affordance 预测它应该成功），系统无法自动检测并尝试替代方案。', font_size=12)
add_dash_item('---对 Value Function 质量的依赖：', '如果 affordance 函数不准确（训练数据不足或状态分布偏移），SayCan 就会错误地排除可行技能或选中不可行技能。', font_size=12)

add_para('')
add_tech_subtitle('展望与后续影响', font_size=12)
add_tech_body('SayCan 是"LLM for Robotics"这一研究方向的开创性工作，其影响力延续至今：（1）直接后续工作 Inner Monologue 为 SayCan 增加了闭环反馈机制，弥补了开环规划的局限；（2）随着 GPT-4、Gemini 等更强模型发布，基于 SayCan 范式的系统将获得更强的语义理解能力；（3）方法论天然适用于任何有明确技能库的场景——工厂车间、医院、家庭服务等；（4）affordance 过滤机制天然为 LLM 的危险建议提供了一道安全屏障——即使 LLM"胡说八道"，只要该行动在 affordance 中被标记为不可行，就不会被执行。SayCan 揭示了一条通向通用具身智能的可能路径：将语义理解、物理可行性和运动控制分层解耦，让每一层专注于自己擅长的子问题，通过简洁的接口（自然语言+概率）互相协作。', font_size=12)

# Save
output_path = r'E:\BaiduSyncdisk\研学习\组会\文献分享\SayCan分享报道.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
